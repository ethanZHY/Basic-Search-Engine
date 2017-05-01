'''
Snippet generation Main process:
reads the query results, for each result document, generate the snippet,
then write the result document and snippet into docx file.
snippet are written in relative order they appears in documents, and query terms in snippet is highlighted by bold
'''

from src.SnipptGenerationModel import SnippetGenerator
from docx import *
import re


sg = SnippetGenerator()
doc = Document()            # create a Document object used for writing into docx file
doc.add_heading('BM25 baseline result with snippet', 0)
# result come from BM25-baseline run:
fr = open('../result/baseline/BM25 baseline result.txt', 'r', encoding='utf-8')
fw = open('../result/BM25 baseline result with snippet.txt', 'w', encoding='utf-8')
fr_query = open('../cacm_query_new.txt', 'r', encoding='utf-8')
query_num = 0
query_num_set = set()
line = fr.readline()
file_name = line[line.find("'")+1:line.rfind(".Txt")]
file_name = file_name.upper() + '.html'
query_number = int(line[:line.find(' ')])
while line != '':
    query_num += 1
    query = fr_query.readline()
    query = sg.remove_punctuation(query).casefold()        # casefold the query, and remove punctuation
    print(query)
    p = doc.add_paragraph('query:' + query)
    while query_number == query_num:
        print(file_name)
        p.add_run(file_name + '\n').bold= True
        sentence_list = sg.doc_process('../cacm/' + file_name)
        terms = sg.find_significant_term(sentence_list)
        labels = sg.rank_significant_sentence(terms, sentence_list)
        query_words = query.split(' ')
        labels = sg.snippet_generation(query_words, labels, sentence_list)
        labels = sorted(labels)     # sort the labels, print the snippet in paragraph going order
        rank = 0
        output = ' '
        for label in labels:
            rank += 1
            sentence = sg.remove_punctuation(sentence_list[label])
            words = sentence.split(' ')
            # new_sentence = ''
            for word in words:
                word = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\u10000-\u10FFFF]+', '', word)
                # remove the ,. following with word and casefold word
                word_without_punc = sg.remove_punctuation(word).casefold()
                # if the word is in query, highlight it
                if word_without_punc in query_words:
                    p.add_run(str(word) + ' ').bold = True
                else:
                    p.add_run(str(word) + ' ')
            p.add_run('....\n')
            if rank == 3:
                break
        line = fr.readline()
        if line == '':
            break
        file_name = line[line.find("'") + 1:line.rfind(".Txt")]
        file_name = file_name.upper() + '.html'
        query_number = int(line[:line.find(' ')])

doc.save('../BM25 baseline result with snippet.docx')
